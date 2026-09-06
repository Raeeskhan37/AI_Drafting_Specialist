import os
import io
import re
import json
import hashlib
import textwrap
from datetime import datetime

import streamlit as st


# ============================================================
# OPTIONAL PACKAGES
# ============================================================

try:
    from groq import Groq
except Exception:
    Groq = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except Exception:
    Document = None

try:
    from fpdf import FPDF
except Exception:
    FPDF = None

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None


# ============================================================
# PAGE CONFIG
# ============================================================

st.set_page_config(
    page_title="DraftForge — AI Document Composer",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ============================================================
# CONSTANTS
# ============================================================

PROFILE_FILE = "user_profile.json"
HISTORY_FILE = "draftforge_history.json"

GROQ_MODEL = "openai/gpt-oss-120b"
WHISPER_MODEL = "whisper-large-v3"

DOCUMENT_TYPES = [
    "Email",
    "Letter",
    "Inquiry",
]

INQUIRY_TYPES = [
    "E&D Inquiry",
    "FFI Inquiry",
]

ED_INDEXES = [
    "Inquiry Reference No.",
    "Subject",
    "Brief of the Inquiry",
    "Articles of Charge / Allegations",
    "Statement of the Accused",
    "Questions / Answers with the Accused",
    "Statements of Witnesses / Officials",
    "Documentary Evidence / Record Examined",
    "Defence / Written Explanation",
    "Findings",
    "Findings on Each Charge",
    "Discussion / Analysis",
    "Conclusion",
    "Recommendations",
    "Documents Recorded",
    "Inquiry Committee",
]

DOCUMENTS_RECORDED = [
    "CNICF",
    "Birth Certificate (BC)",
    "Marriage Certificate",
    "CNICs",
    "Domicile",
    "Affidavit",
    "Complaint / Application",
    "Written Explanation",
    "Official Record",
    "Statement of Accused",
    "Statement of Witness",
    "Other",
]

COMMITTEE_ROLES = [
    "Convener of Inquiry",
    "Member 1",
    "Member 2",
    "Departmental Representative",
]


# ============================================================
# DEFAULT SESSION STATE
# ============================================================

DEFAULTS = {
    "document_type": "Email",
    "inquiry_type": "E&D Inquiry",
    "index_data": [],
    "generated_draft": "",
    "editable_draft": "",
    "document_editor": "",
    "editor_sync": "",
    "edit_instruction": "",
    "edit_instruction_sync": "",
    "email_instruction": "",
    "letter_instruction": "",
    "email_voice_hash": "",
    "letter_voice_hash": "",
    "profile": {},
    "history": [],
    "show_history": False,
    "show_profile": False,
    "documents_version": 0,
}

for key, value in DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = value


# ============================================================
# CUSTOM CSS — MODERN DRAFTFORGE UI
# ============================================================

st.markdown(
    """
<style>

/* ==========================================================
   GLOBAL
   ========================================================== */

.stApp {
    background:
        radial-gradient(circle at 10% 0%, rgba(37, 99, 235, 0.07), transparent 28%),
        radial-gradient(circle at 90% 5%, rgba(124, 58, 237, 0.06), transparent 25%),
        #f6f8fc;
}

.block-container {
    max-width: 1450px;
    padding-top: 1.2rem;
    padding-bottom: 4rem;
}

/* Hide default Streamlit chrome */
#MainMenu {
    visibility: hidden;
}

footer {
    visibility: hidden;
}

header {
    background: transparent !important;
}


/* ==========================================================
   TYPOGRAPHY
   ========================================================== */

h1, h2, h3, h4 {
    color: #172033 !important;
}

p, label, span {
    color: #334155;
}


/* ==========================================================
   SIDEBAR
   ========================================================== */

section[data-testid="stSidebar"] {
    background: linear-gradient(
        180deg,
        #101828 0%,
        #172554 48%,
        #111827 100%
    );
    border-right: 1px solid rgba(255,255,255,0.08);
}

section[data-testid="stSidebar"] * {
    color: #e5e7eb !important;
}

section[data-testid="stSidebar"] .stButton button {
    background: transparent !important;
    border: 1px solid transparent !important;
    color: #e5e7eb !important;
    text-align: left !important;
    justify-content: flex-start !important;
}

section[data-testid="stSidebar"] .stButton button:hover {
    background: rgba(255,255,255,0.08) !important;
    border-color: rgba(255,255,255,0.12) !important;
}

.sidebar-brand {
    padding: 10px 4px 24px 4px;
}

.sidebar-logo {
    width: 48px;
    height: 48px;
    border-radius: 15px;
    background: linear-gradient(135deg, #60a5fa, #7c3aed);
    display: flex;
    align-items: center;
    justify-content: center;
    color: white !important;
    font-size: 25px;
    font-weight: 800;
    box-shadow: 0 10px 30px rgba(0,0,0,0.25);
}

.sidebar-title {
    font-size: 22px;
    font-weight: 800;
    color: white !important;
    margin-top: 12px;
}

.sidebar-subtitle {
    font-size: 12px;
    color: #94a3b8 !important;
    margin-top: -3px;
}

.sidebar-section {
    color: #94a3b8 !important;
    font-size: 11px;
    font-weight: 800;
    letter-spacing: 1.2px;
    text-transform: uppercase;
    margin: 22px 4px 7px;
}


/* ==========================================================
   HERO
   ========================================================== */

.hero {
    background:
        linear-gradient(
            135deg,
            rgba(15, 23, 42, 0.98),
            rgba(30, 64, 175, 0.96) 58%,
            rgba(79, 70, 229, 0.94)
        );
    border-radius: 25px;
    padding: 30px 34px;
    color: white;
    box-shadow: 0 18px 50px rgba(15, 23, 42, 0.18);
    position: relative;
    overflow: hidden;
    margin-bottom: 25px;
}

.hero:after {
    content: "";
    position: absolute;
    width: 280px;
    height: 280px;
    right: -80px;
    top: -120px;
    border-radius: 50%;
    background: rgba(255,255,255,0.08);
}

.hero-title {
    color: white !important;
    font-size: 38px;
    font-weight: 850;
    letter-spacing: -1px;
    margin: 0;
}

.hero-subtitle {
    color: #dbeafe !important;
    font-size: 15px;
    margin-top: 6px;
    max-width: 780px;
}

.hero-badge {
    display: inline-block;
    margin-top: 17px;
    padding: 7px 13px;
    border-radius: 999px;
    background: rgba(255,255,255,0.12);
    border: 1px solid rgba(255,255,255,0.18);
    color: #e0e7ff !important;
    font-size: 12px;
    font-weight: 700;
}


/* ==========================================================
   WORKFLOW
   ========================================================== */

.workflow {
    display: flex;
    align-items: center;
    gap: 10px;
    margin: 5px 0 22px;
}

.workflow-step {
    flex: 1;
    padding: 12px 15px;
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 14px;
    box-shadow: 0 4px 14px rgba(15,23,42,0.04);
}

.workflow-number {
    display: inline-flex;
    width: 27px;
    height: 27px;
    border-radius: 50%;
    align-items: center;
    justify-content: center;
    background: #eef2ff;
    color: #4338ca !important;
    font-weight: 800;
    margin-right: 8px;
}

.workflow-text {
    font-weight: 700;
    color: #1e293b !important;
}


/* ==========================================================
   CARDS
   ========================================================== */

.section-card {
    background: rgba(255,255,255,0.94);
    border: 1px solid #e2e8f0;
    border-radius: 21px;
    padding: 22px;
    box-shadow: 0 8px 30px rgba(15,23,42,0.055);
    margin-bottom: 18px;
}

.section-title {
    font-size: 19px;
    font-weight: 800;
    color: #172033 !important;
    margin-bottom: 3px;
}

.section-description {
    color: #64748b !important;
    font-size: 13px;
    margin-bottom: 18px;
}


/* ==========================================================
   DOCUMENT CARDS
   ========================================================== */

.doc-card {
    border: 2px solid #e2e8f0;
    border-radius: 18px;
    padding: 17px;
    min-height: 125px;
    background: white;
    box-shadow: 0 5px 18px rgba(15,23,42,0.04);
}

.doc-card-icon {
    font-size: 27px;
}

.doc-card-title {
    font-size: 17px;
    font-weight: 800;
    color: #172033 !important;
    margin-top: 5px;
}

.doc-card-text {
    font-size: 12px;
    color: #64748b !important;
    line-height: 1.45;
}


/* ==========================================================
   INPUTS
   ========================================================== */

div[data-testid="stTextInput"] input,
div[data-testid="stTextArea"] textarea {
    background: #ffffff !important;
    color: #111827 !important;
    border: 1.5px solid #cbd5e1 !important;
    border-radius: 12px !important;
    box-shadow: none !important;
}

div[data-testid="stTextInput"] input:focus,
div[data-testid="stTextArea"] textarea:focus {
    border-color: #4f46e5 !important;
    box-shadow: 0 0 0 3px rgba(79,70,229,0.10) !important;
}

div[data-testid="stTextInput"] label,
div[data-testid="stTextArea"] label,
div[data-testid="stSelectbox"] label,
div[data-testid="stMultiSelect"] label {
    color: #334155 !important;
    font-weight: 700 !important;
}

div[data-testid="stTextArea"] textarea {
    line-height: 1.55 !important;
}


/* ==========================================================
   SELECTBOX
   ========================================================== */

div[data-testid="stSelectbox"] > div > div {
    background: white !important;
    border-radius: 12px !important;
}


/* ==========================================================
   BUTTONS
   ========================================================== */

.stButton button {
    border-radius: 11px !important;
    border: 1px solid #cbd5e1 !important;
    font-weight: 700 !important;
    min-height: 42px !important;
    transition: all 0.15s ease-in-out !important;
}

.stButton button:hover {
    transform: translateY(-1px);
    box-shadow: 0 7px 18px rgba(15,23,42,0.10);
}

.generate-button button {
    min-height: 54px !important;
    border: none !important;
    color: white !important;
    font-size: 16px !important;
    background: linear-gradient(135deg, #2563eb, #4f46e5) !important;
    box-shadow: 0 10px 25px rgba(37,99,235,0.25);
}

.generate-button button:hover {
    box-shadow: 0 14px 30px rgba(37,99,235,0.32);
}


/* ==========================================================
   VOICE CARD
   ========================================================== */

.voice-card {
    background: linear-gradient(135deg, #f8fafc, #eef2ff);
    border: 1px solid #dbeafe;
    border-radius: 16px;
    padding: 14px 16px;
    margin: 8px 0 12px;
}

.voice-title {
    font-weight: 800;
    color: #1e293b !important;
}

.voice-description {
    color: #64748b !important;
    font-size: 12px;
}


/* ==========================================================
   INQUIRY INDEX
   ========================================================== */

.index-card {
    background: white;
    border: 1px solid #dbe3ef;
    border-radius: 17px;
    padding: 18px;
    margin-bottom: 15px;
    box-shadow: 0 5px 20px rgba(15,23,42,0.04);
}

.index-header {
    display: flex;
    align-items: center;
    gap: 10px;
    margin-bottom: 13px;
}

.index-number {
    width: 31px;
    height: 31px;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    border-radius: 10px;
    background: #eef2ff;
    color: #4338ca !important;
    font-weight: 850;
    font-size: 13px;
}

.index-name {
    font-weight: 800;
    color: #172033 !important;
}


/* ==========================================================
   COMMITTEE
   ========================================================== */

.committee-role {
    background: linear-gradient(135deg, #eef2ff, #f8fafc);
    border: 1px solid #c7d2fe;
    border-radius: 15px;
    padding: 12px 15px;
    margin: 7px 0 12px;
}

.committee-role-title {
    color: #312e81 !important;
    font-weight: 850;
    font-size: 15px;
}

.committee-role-subtitle {
    color: #64748b !important;
    font-size: 11px;
}


/* ==========================================================
   SELECTED PANEL
   ========================================================== */

.selected-panel {
    background: #ffffff;
    border: 1px solid #dbe3ef;
    border-radius: 18px;
    padding: 17px;
    margin-bottom: 18px;
}

.selected-count {
    font-size: 25px;
    font-weight: 850;
    color: #4338ca !important;
}

.selected-label {
    color: #64748b !important;
    font-size: 12px;
}


/* ==========================================================
   EDITOR
   ========================================================== */

.editor-card {
    background: white;
    border: 1px solid #dbe3ef;
    border-radius: 20px;
    padding: 20px;
    box-shadow: 0 8px 30px rgba(15,23,42,0.05);
}

.editor-heading {
    font-size: 21px;
    font-weight: 850;
    color: #172033 !important;
}


/* ==========================================================
   EXPORT
   ========================================================== */

.export-card {
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 18px;
    padding: 17px;
    margin-top: 15px;
}


/* ==========================================================
   INFO / WARNING
   ========================================================== */

div[data-testid="stAlert"] {
    border-radius: 13px !important;
}


/* ==========================================================
   DIVIDER
   ========================================================== */

hr {
    border-color: #e2e8f0 !important;
}


/* ==========================================================
   FOOTER
   ========================================================== */

.app-footer {
    text-align: center;
    padding: 30px 0 10px;
    color: #94a3b8 !important;
    font-size: 12px;
}

</style>
""",
    unsafe_allow_html=True,
)


# ============================================================
# UTILITY FUNCTIONS
# ============================================================

def load_json(filename, default):
    try:
        if os.path.exists(filename):
            with open(filename, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return default


def save_json(filename, data):
    try:
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception:
        return False


def clean_ai_text(text):
    if not text:
        return ""

    text = str(text).strip()

    text = re.sub(
        r"^```(?:text|markdown)?\s*",
        "",
        text,
        flags=re.IGNORECASE,
    )

    text = re.sub(r"\s*```$", "", text)

    return text.strip()


def profile_signature(profile):
    if not profile:
        return ""

    name = str(profile.get("Name", "")).strip()
    designation = str(profile.get("Designation", "")).strip()
    contact = str(profile.get("Contact No.", "")).strip()
    station = str(profile.get("Current Station", "")).strip()

    return "\n".join(
        [
            name,
            designation,
            f"Contact No.: {contact}" if contact else "",
            f"Current Station: {station}" if station else "",
        ]
    ).strip()


def get_groq_client():
    if Groq is None:
        return None

    api_key = ""

    try:
        api_key = st.secrets.get("GROQ_API_KEY", "")
    except Exception:
        pass

    if not api_key:
        api_key = os.getenv("GROQ_API_KEY", "")

    if not api_key:
        return None

    try:
        return Groq(api_key=api_key)
    except Exception:
        return None


# ============================================================
# PROFILE / HISTORY
# ============================================================

if not st.session_state.profile:
    st.session_state.profile = load_json(PROFILE_FILE, {})

if not st.session_state.history:
    st.session_state.history = load_json(HISTORY_FILE, [])


# ============================================================
# AI DOCUMENT GENERATION
# ============================================================

def generate_ai_document(document_type, instruction, profile):
    client = get_groq_client()

    if client is None:
        return (
            "Groq API is not configured. Please add GROQ_API_KEY to "
            "Streamlit Secrets or the environment variables."
        )

    signature = profile_signature(profile)

    if document_type == "Email":
        system_prompt = """
You are DraftForge, an AI assistant for drafting professional official emails.

Write a polished official email from the user's natural-language instructions.

Rules:
- Correct spelling, grammar, punctuation and obvious voice-transcription errors.
- Preserve the user's intended meaning.
- Do not invent facts.
- Do not invent names, dates, reference numbers, allegations, evidence,
  events or commitments.
- Use professional official English.
- Do not add an unsupported subject.
- Do not add a sender signature because the application adds the profile
  automatically.
- Return only the email body.
"""

    elif document_type == "Letter":
        system_prompt = """
You are DraftForge, an AI assistant for drafting professional official letters.

Write a polished official letter from the user's natural-language instructions.

Rules:
- Correct spelling, grammar, punctuation and obvious voice-transcription errors.
- Preserve the user's intended meaning.
- Do not invent facts.
- Do not invent names, dates, reference numbers, allegations, evidence,
  events or commitments.
- Use professional official English.
- Do not add a sender signature because the application adds the profile
  automatically.
- Return only the letter content.
"""

    else:
        system_prompt = """
You are DraftForge, an AI assistant for official inquiry documentation.

Draft the requested inquiry content professionally.

Rules:
- Preserve the supplied facts exactly.
- Do not invent allegations, witnesses, evidence, findings,
  recommendations, dates, names or events.
- Do not introduce sections that were not requested.
- Correct grammar, spelling and obvious voice transcription errors.
- Use formal official English.
- Return only the requested content.
"""

    user_prompt = f"""
Document type:
{document_type}

User instructions:
{instruction}

User profile:
{signature}
"""

    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
        )

        return clean_ai_text(response.choices[0].message.content)

    except Exception as e:
        return f"AI generation failed: {e}"


# ============================================================
# VOICE TRANSCRIPTION
# ============================================================

def transcribe_audio(audio_bytes):
    client = get_groq_client()

    if client is None:
        return "Groq API is not configured."

    try:
        audio_file = io.BytesIO(audio_bytes)
        audio_file.name = "voice_input.wav"

        transcription = client.audio.transcriptions.create(
            file=audio_file,
            model=WHISPER_MODEL,
        )

        return str(transcription.text).strip()

    except Exception as e:
        return f"Voice transcription failed: {e}"


def process_voice_input(audio_value, state_key):
    if audio_value is None:
        return ""

    try:
        audio_bytes = audio_value.getvalue()
    except Exception:
        try:
            audio_bytes = audio_value.read()
        except Exception:
            return ""

    if not audio_bytes:
        return ""

    audio_hash = hashlib.sha256(audio_bytes).hexdigest()

    if st.session_state.get(state_key) == audio_hash:
        return ""

    transcript = transcribe_audio(audio_bytes)

    if transcript.startswith("Voice transcription failed"):
        st.error(transcript)
        return ""

    if transcript == "Groq API is not configured.":
        st.error(transcript)
        return ""

    st.session_state[state_key] = audio_hash

    return transcript


# ============================================================
# DOCUMENT HELPERS
# ============================================================

def documents_text(documents):
    if not documents:
        return "No information was provided for this index."

    lines = []

    for i, document in enumerate(documents, start=1):
        lines.append(f"Annex-{chr(64 + i)} — {document}")

    return "\n".join(lines)


def normalize_committee(committee):
    result = {
        role: {
            "ERP#": "",
            "Name": "",
            "Designation": "",
        }
        for role in COMMITTEE_ROLES
    }

    if isinstance(committee, dict):
        for role in COMMITTEE_ROLES:
            current = committee.get(role, {})

            if isinstance(current, dict):
                result[role] = {
                    "ERP#": str(
                        current.get("ERP#", current.get("erp", "")) or ""
                    ),
                    "Name": str(
                        current.get("Name", current.get("name", "")) or ""
                    ),
                    "Designation": str(
                        current.get(
                            "Designation",
                            current.get("designation", ""),
                        )
                        or ""
                    ),
                }

    elif isinstance(committee, list):
        for member in committee:
            if not isinstance(member, dict):
                continue

            role = member.get("role", "")

            if role in result:
                result[role] = {
                    "ERP#": str(
                        member.get("ERP#", member.get("erp", "")) or ""
                    ),
                    "Name": str(
                        member.get("Name", member.get("name", "")) or ""
                    ),
                    "Designation": str(
                        member.get(
                            "Designation",
                            member.get("designation", ""),
                        )
                        or ""
                    ),
                }

    return result


def committee_text(committee):
    committee = normalize_committee(committee)

    result = []

    for role in COMMITTEE_ROLES:
        member = committee.get(role, {})

        erp = str(member.get("ERP#", "") or "").strip()
        name = str(member.get("Name", "") or "").strip()
        designation = str(member.get("Designation", "") or "").strip()

        if not (erp or name or designation):
            continue

        parts = []

        if erp:
            parts.append(f"ERP#: {erp}")

        if name:
            parts.append(f"Name: {name}")

        if designation:
            parts.append(f"Designation: {designation}")

        result.append(f"{role} — " + ", ".join(parts))

    if not result:
        return "No information was provided for this index."

    return "\n".join(result)


def qa_markdown(rows):
    if not rows:
        return "No information was provided for this index."

    lines = [
        "| Questions | Answers |",
        "|---|---|",
    ]

    for i, row in enumerate(rows, start=1):
        if isinstance(row, dict):
            question = str(row.get("question", "")).strip()
            answer = str(row.get("answer", "")).strip()
        else:
            question = ""
            answer = ""

        question = question.replace("|", "\\|").replace("\n", " ")
        answer = answer.replace("|", "\\|").replace("\n", " ")

        lines.append(f"| {i}. {question} | {answer} |")

    return "\n".join(lines)


# ============================================================
# INQUIRY REPORT
# ============================================================

def generate_inquiry_report(index_data, inquiry_type):
    if not index_data:
        return "No inquiry sections were selected."

    output = []

    output.append(inquiry_type.upper())
    output.append("")

    for position, item in enumerate(index_data, start=1):
        title = item.get("title", "")
        base_name = item.get("base_name", title)

        output.append(f"{position}. {title}")
        output.append("")

        if base_name == "Documents Recorded":
            documents = item.get("documents", [])

            if documents:
                for i, document in enumerate(documents):
                    letter = chr(65 + i)
                    output.append(
                        f"Annex-{letter}: {document}"
                    )
            else:
                output.append(
                    "No information was provided for this index."
                )

        elif base_name == "Inquiry Committee":
            committee = committee_text(
                item.get("committee", {})
            )

            output.append(committee)

        elif base_name == "Questions / Answers with the Accused":
            rows = item.get("qa", [])

            output.append(qa_markdown(rows))

        else:
            content = str(item.get("content", "")).strip()

            if content:
                output.append(content)
            else:
                output.append(
                    "No information was provided for this index."
                )

        output.append("")

    return "\n".join(output).strip()


# ============================================================
# AI EDITING
# ============================================================

def ai_edit_document(document, instruction):
    client = get_groq_client()

    if client is None:
        return "Groq API is not configured."

    prompt = f"""
You are editing an official document.

Original document:
{document}

Requested editing instruction:
{instruction}

Rules:
- Make only the requested changes.
- Preserve the original meaning.
- Do not invent facts.
- Do not add unsupported names, dates, allegations, evidence,
  findings, recommendations or events.
- Keep the document professional and official.
- Return only the revised document.
"""

    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a precise official-document editor."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            temperature=0.15,
        )

        return clean_ai_text(response.choices[0].message.content)

    except Exception as e:
        return f"AI editing failed: {e}"


# ============================================================
# HISTORY
# ============================================================

def save_history(title, document):
    history = st.session_state.history

    history.insert(
        0,
        {
            "title": title,
            "document": document,
            "created": datetime.now().strftime(
                "%Y-%m-%d %H:%M"
            ),
        },
    )

    st.session_state.history = history[:30]

    save_json(HISTORY_FILE, st.session_state.history)


# ============================================================
# PDF HELPERS
# ============================================================

def safe_pdf_text(text):
    replacements = {
        "—": "-",
        "–": "-",
        "•": "-",
        "’": "'",
        "‘": "'",
        "“": '"',
        "”": '"',
        "…": "...",
        "✓": "[OK]",
        "✦": "*",
        "🎤": "[Voice]",
        "👤": "",
        "📄": "",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text.encode(
        "latin-1",
        errors="replace",
    ).decode("latin-1")


def create_pdf(text):
    if FPDF is None:
        return None

    pdf = FPDF()
    pdf.set_auto_page_break(
        auto=True,
        margin=18,
    )
    pdf.add_page()

    pdf.set_left_margin(18)
    pdf.set_right_margin(18)

    pdf.set_font(
        "Arial",
        size=11,
    )

    usable_width = (
        pdf.w - pdf.l_margin - pdf.r_margin
    )

    for raw_line in text.splitlines():
        line = safe_pdf_text(raw_line)

        if not line.strip():
            pdf.ln(5)
            continue

        wrapped = textwrap.wrap(
            line,
            width=90,
            break_long_words=True,
            break_on_hyphens=True,
        )

        if not wrapped:
            pdf.ln(5)
            continue

        for wrapped_line in wrapped:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(
                usable_width,
                6,
                wrapped_line,
            )

    return bytes(pdf.output(dest="S"))


# ============================================================
# DOCX EXPORT
# ============================================================

def create_docx(text):
    if Document is None:
        return None

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = text.splitlines()

    for line in lines:
        if line.startswith("| Questions | Answers |"):
            continue

        if line.startswith("|---"):
            continue

        if line.startswith("| ") and line.endswith(" |"):
            cells = [
                c.strip()
                for c in line.strip("|").split("|")
            ]

            if len(cells) == 2:
                if not hasattr(create_docx, "_table"):
                    pass

                continue

        p = document.add_paragraph()

        if re.match(r"^\d+\.", line):
            run = p.add_run(line)
            run.bold = True
        else:
            p.add_run(line)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output.getvalue()


# ============================================================
# PNG EXPORT
# ============================================================

def create_png(text):
    if Image is None:
        return None

    width = 1600
    margin = 70
    line_height = 30

    font = None

    try:
        font = ImageFont.truetype(
            "DejaVuSans.ttf",
            22,
        )
    except Exception:
        try:
            font = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                22,
            )
        except Exception:
            font = ImageFont.load_default()

    wrapped_lines = []

    for raw_line in text.splitlines():
        if not raw_line:
            wrapped_lines.append("")
            continue

        wrapped_lines.extend(
            textwrap.wrap(
                raw_line,
                width=105,
                break_long_words=True,
            )
        )

    height = max(
        500,
        margin * 2 + len(wrapped_lines) * line_height,
    )

    image = Image.new(
        "RGB",
        (width, height),
        "white",
    )

    draw = ImageDraw.Draw(image)

    y = margin

    for line in wrapped_lines:
        draw.text(
            (margin, y),
            line,
            fill="black",
            font=font,
        )
        y += line_height

    output = io.BytesIO()
    image.save(
        output,
        format="PNG",
    )
    output.seek(0)

    return output.getvalue()


# ============================================================
# SIDEBAR
# ============================================================

with st.sidebar:

    st.markdown(
        """
        <div class="sidebar-brand">
            <div class="sidebar-logo">✦</div>
            <div class="sidebar-title">DraftForge</div>
            <div class="sidebar-subtitle">
                AI Document Composer
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        '<div class="sidebar-section">Workspace</div>',
        unsafe_allow_html=True,
    )

    if st.button("✦  New Document", use_container_width=True):
        st.session_state.generated_draft = ""
        st.session_state.editable_draft = ""
        st.session_state.document_editor = ""
        st.session_state.editor_sync = ""
        st.session_state.index_data = []
        st.session_state.email_instruction = ""
        st.session_state.letter_instruction = ""
        st.session_state.show_history = False
        st.rerun()

    if st.button("🗂️  My Documents", use_container_width=True):
        st.session_state.show_history = True
        st.session_state.show_profile = False

    if st.button("👤  My Profile", use_container_width=True):
        st.session_state.show_profile = True
        st.session_state.show_history = False

    st.markdown(
        '<div class="sidebar-section">Help</div>',
        unsafe_allow_html=True,
    )

    with st.expander("💡 Tips & Templates"):

        st.markdown(
            """
            **Write naturally**  
            DraftForge converts your instructions into professional
            official English.

            **🎤 Speak instead of typing**  
            Voice and typed information can be combined.

            **📄 Inquiry sections**  
            Add sections in any order.

            **🔁 Repeat sections**  
            The same inquiry index can be added multiple times.
            """
        )

    with st.expander("ℹ️ About DraftForge"):

        st.markdown(
            """
            **DraftForge** is an AI-assisted drafting workspace for
            professional official correspondence and inquiry
            documentation.

            **Supported documents**

            • Email  
            • Letter  
            • E&D Inquiry  
            • FFI Inquiry
            """
        )

    st.markdown("---")

    st.caption("DraftForge — AI Document Composer")
    st.caption("Developed by: Raees Khan — Assistant Director, NADRA")


# ============================================================
# PROFILE PANEL
# ============================================================

if st.session_state.show_profile:

    st.markdown(
        """
        <div class="section-card">
            <div class="section-title">👤 My Profile</div>
            <div class="section-description">
                Your profile is automatically used when preparing
                your document signature.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    profile = st.session_state.profile

    with st.form("profile_form"):

        name = st.text_input(
            "Name",
            value=profile.get("Name", ""),
        )

        designation = st.text_input(
            "Designation",
            value=profile.get("Designation", ""),
        )

        contact = st.text_input(
            "Contact No.",
            value=profile.get("Contact No.", ""),
        )

        station = st.text_input(
            "Current Station",
            value=profile.get("Current Station", ""),
        )

        submitted = st.form_submit_button(
            "💾 Save Profile",
            use_container_width=True,
        )

        if submitted:

            st.session_state.profile = {
                "Name": name.strip(),
                "Designation": designation.strip(),
                "Contact No.": contact.strip(),
                "Current Station": station.strip(),
            }

            save_json(
                PROFILE_FILE,
                st.session_state.profile,
            )

            st.success("Profile saved successfully.")

    st.stop()


# ============================================================
# HISTORY PANEL
# ============================================================

if st.session_state.show_history:

    st.markdown(
        """
        <div class="section-card">
            <div class="section-title">🗂️ My Documents</div>
            <div class="section-description">
                Previously generated documents saved by DraftForge.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    history = st.session_state.history

    if not history:
        st.info(
            "No saved documents yet. Generate a document and it "
            "will appear here."
        )

    for i, item in enumerate(history):

        title = item.get(
            "title",
            "Untitled Document",
        )

        created = item.get(
            "created",
            "",
        )

        with st.expander(
            f"📄 {title}  ·  {created}"
        ):

            st.text_area(
                "Document",
                value=item.get("document", ""),
                height=250,
                key=f"history_view_{i}",
            )

            if st.button(
                "Open in Editor",
                key=f"history_open_{i}",
            ):

                document = item.get(
                    "document",
                    "",
                )

                st.session_state.generated_draft = document
                st.session_state.editable_draft = document
                st.session_state.editor_sync = document
                st.session_state.show_history = False

                st.rerun()

    st.stop()


# ============================================================
# HERO
# ============================================================

st.markdown(
    """
    <div class="hero">
        <div class="hero-title">✦ DraftForge</div>
        <div class="hero-subtitle">
            Transform natural-language instructions into polished,
            professional official documents — faster and with less effort.
        </div>
        <div class="hero-badge">
            AI-ASSISTED OFFICIAL DOCUMENT WORKSPACE
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# WORKFLOW
# ============================================================

st.markdown(
    """
    <div class="workflow">
        <div class="workflow-step">
            <span class="workflow-number">1</span>
            <span class="workflow-text">Choose Document</span>
        </div>

        <div class="workflow-step">
            <span class="workflow-number">2</span>
            <span class="workflow-text">Provide Information</span>
        </div>

        <div class="workflow-step">
            <span class="workflow-number">3</span>
            <span class="workflow-text">Generate & Export</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# STEP 1 — DOCUMENT TYPE
# ============================================================

st.markdown(
    """
    <div class="section-card">
        <div class="section-title">
            ① Choose your document
        </div>
        <div class="section-description">
            Select what you want DraftForge to prepare.
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

doc_cols = st.columns(3)

document_descriptions = {
    "Email": (
        "📧",
        "Professional official email",
    ),
    "Letter": (
        "📄",
        "Formal official correspondence",
    ),
    "Inquiry": (
        "⚖️",
        "Structured inquiry documentation",
    ),
}

for col, doc_type in zip(
    doc_cols,
    DOCUMENT_TYPES,
):

    icon, description = document_descriptions[
        doc_type
    ]

    with col:

        st.markdown(
            f"""
            <div class="doc-card">
                <div class="doc-card-icon">{icon}</div>
                <div class="doc-card-title">{doc_type}</div>
                <div class="doc-card-text">
                    {description}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        selected = st.button(
            (
                "✓ Selected"
                if st.session_state.document_type == doc_type
                else f"Use {doc_type}"
            ),
            key=f"select_doc_{doc_type}",
            use_container_width=True,
        )

        if selected:
            st.session_state.document_type = doc_type
            st.session_state.generated_draft = ""
            st.session_state.editable_draft = ""
            st.session_state.document_editor = ""
            st.session_state.editor_sync = ""

            if doc_type == "Inquiry":
                st.session_state.index_data = []

            st.rerun()


# ============================================================
# CURRENT DOCUMENT INDICATOR
# ============================================================

current_type = st.session_state.document_type

st.markdown(
    f"""
    <div style="
        margin: 12px 0 20px;
        padding: 10px 15px;
        border-radius: 12px;
        background: #eef2ff;
        border: 1px solid #c7d2fe;
        color: #3730a3;
        font-weight: 750;
        font-size: 13px;
    ">
        CURRENT DOCUMENT: {current_type.upper()}
    </div>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# STEP 2 — EMAIL
# ============================================================

if current_type == "Email":

    st.markdown(
        """
        <div class="section-card">
            <div class="section-title">
                ② Provide Information
            </div>
            <div class="section-description">
                Tell DraftForge what you want to communicate.
                You can type, speak, or combine both.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="voice-card">
            <div class="voice-title">
                🎤 Voice + Text Input
            </div>
            <div class="voice-description">
                Record your instructions. Your transcription will be
                added to the same text box below.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # IMPORTANT:
    # Voice is processed BEFORE text_area is instantiated.
    audio = st.audio_input(
        "🎤 Record your instructions",
        key="email_audio",
    )

    transcript = process_voice_input(
        audio,
        "email_voice_hash",
    )

    if transcript:
        existing = st.session_state.get(
            "email_instruction",
            "",
        ).strip()

        if existing:
            st.session_state.email_instruction = (
                existing + "\n" + transcript
            )
        else:
            st.session_state.email_instruction = transcript

    st.text_area(
        "What should the email say?",
        key="email_instruction",
        height=230,
        placeholder=(
            "Example: Please write an email to the concerned "
            "office informing them that the report has been "
            "completed and requesting them to review it."
        ),
    )

    st.caption(
        "💡 You can type normally, use the microphone, or use both."
    )

    st.markdown("<br>", unsafe_allow_html=True)

    generate_email = st.button(
        "✦ Generate Professional Email",
        use_container_width=True,
        type="primary",
    )

    if generate_email:

        instruction = st.session_state.email_instruction.strip()

        if not instruction:
            st.warning(
                "Please provide some information by typing or speaking."
            )
        else:

            with st.spinner(
                "DraftForge is preparing your email..."
            ):

                final_document = generate_ai_document(
                    "Email",
                    instruction,
                    st.session_state.profile,
                )

            st.session_state.generated_draft = final_document
            st.session_state.editable_draft = final_document
            st.session_state.editor_sync = final_document

            save_history(
                "Email",
                final_document,
            )

            st.rerun()


# ============================================================
# STEP 2 — LETTER
# ============================================================

elif current_type == "Letter":

    st.markdown(
        """
        <div class="section-card">
            <div class="section-title">
                ② Provide Information
            </div>
            <div class="section-description">
                Describe the purpose and contents of your letter.
                Speak naturally or type your instructions.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="voice-card">
            <div class="voice-title">
                🎤 Voice + Text Input
            </div>
            <div class="voice-description">
                Your voice transcription and typed information
                are combined in the same input field.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    audio = st.audio_input(
        "🎤 Record your instructions",
        key="letter_audio",
    )

    transcript = process_voice_input(
        audio,
        "letter_voice_hash",
    )

    if transcript:

        existing = st.session_state.get(
            "letter_instruction",
            "",
        ).strip()

        if existing:
            st.session_state.letter_instruction = (
                existing + "\n" + transcript
            )
        else:
            st.session_state.letter_instruction = transcript

    st.text_area(
        "What should the letter say?",
        key="letter_instruction",
        height=230,
        placeholder=(
            "Example: Draft a formal letter regarding the "
            "pending matter and request necessary action."
        ),
    )

    st.caption(
        "💡 Speak naturally. DraftForge will convert your instructions "
        "into formal official English."
    )

    st.markdown("<br>", unsafe_allow_html=True)

    generate_letter = st.button(
        "✦ Generate Professional Letter",
        use_container_width=True,
        type="primary",
    )

    if generate_letter:

        instruction = st.session_state.letter_instruction.strip()

        if not instruction:
            st.warning(
                "Please provide some information by typing or speaking."
            )
        else:

            with st.spinner(
                "DraftForge is preparing your letter..."
            ):

                final_document = generate_ai_document(
                    "Letter",
                    instruction,
                    st.session_state.profile,
                )

            st.session_state.generated_draft = final_document
            st.session_state.editable_draft = final_document
            st.session_state.editor_sync = final_document

            save_history(
                "Letter",
                final_document,
            )

            st.rerun()


# ============================================================
# STEP 2 — INQUIRY
# ============================================================

elif current_type == "Inquiry":

    st.markdown(
        """
        <div class="section-card">
            <div class="section-title">
                ② Build your inquiry
            </div>
            <div class="section-description">
                Choose the inquiry type, then add only the sections
                you actually need. You can add the same section more
                than once.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    inquiry_col1, inquiry_col2 = st.columns(
        [1, 2]
    )

    with inquiry_col1:

        inquiry_type = st.selectbox(
            "Inquiry Type",
            INQUIRY_TYPES,
            index=INQUIRY_TYPES.index(
                st.session_state.inquiry_type
            ),
        )

        st.session_state.inquiry_type = inquiry_type

    with inquiry_col2:

        if inquiry_type == "FFI Inquiry":

            st.info(
                "FFI Inquiry is currently under construction / "
                "under process."
            )

        else:

            st.success(
                "E&D Inquiry — select and arrange the sections "
                "required for your report."
            )


    # --------------------------------------------------------
    # INDEX SELECTION
    # --------------------------------------------------------

    if inquiry_type == "E&D Inquiry":

        st.markdown(
            """
            <div class="section-card">
                <div class="section-title">
                    📑 Add Inquiry Sections
                </div>
                <div class="section-description">
                    Select a section and add it below. Sections are
                    generated in exactly the order you add them.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        available_indexes = ED_INDEXES

        index_choice = st.selectbox(
            "Select an inquiry index to add",
            available_indexes,
            key="new_index_choice",
        )

        add_index = st.button(
            "＋ Add This Section",
            use_container_width=True,
        )

        if add_index:

            base_name = index_choice

            same_count = sum(
                1
                for x in st.session_state.index_data
                if x.get("base_name") == base_name
            )

            if base_name in [
                "Statement of the Accused",
                "Questions / Answers with the Accused",
            ]:

                title = (
                    f"{base_name} "
                    f"No. {same_count + 1}"
                )

            else:

                title = base_name

                if same_count > 0:
                    title = (
                        f"{base_name} "
                        f"No. {same_count + 1}"
                    )

            new_item = {
                "base_name": base_name,
                "title": title,
                "content": "",
                "documents": [],
                "committee": {
                    role: {
                        "ERP#": "",
                        "Name": "",
                        "Designation": "",
                    }
                    for role in COMMITTEE_ROLES
                },
                "qa": [],
            }

            st.session_state.index_data.append(
                new_item
            )

            st.rerun()


        # ----------------------------------------------------
        # SELECTED SECTION SUMMARY
        # ----------------------------------------------------

        selected_count = len(
            st.session_state.index_data
        )

        st.markdown(
            f"""
            <div class="selected-panel">
                <div class="selected-count">
                    {selected_count}
                </div>
                <div class="selected-label">
                    inquiry section{"s" if selected_count != 1 else ""}
                    currently selected
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if selected_count == 0:

            st.info(
                "No sections have been added yet. Choose an index "
                "above and click “Add This Section”."
            )


        # ----------------------------------------------------
        # RENDER EACH SELECTED INDEX
        # ----------------------------------------------------

        for position, item in enumerate(
            st.session_state.index_data
        ):

            base_name = item.get(
                "base_name",
                "",
            )

            title = item.get(
                "title",
                base_name,
            )

            st.markdown(
                f"""
                <div class="index-card">
                    <div class="index-header">
                        <span class="index-number">
                            {position + 1}
                        </span>
                        <span class="index-name">
                            {title}
                        </span>
                    </div>
                """,
                unsafe_allow_html=True,
            )


            # =================================================
            # DOCUMENTS RECORDED
            # =================================================

            if base_name == "Documents Recorded":

                st.markdown(
                    "##### 📎 Select Documents to be Recorded"
                )

                current_documents = item.get(
                    "documents",
                    [],
                )

                selected_documents = st.multiselect(
                    "Available documents",
                    DOCUMENTS_RECORDED,
                    default=current_documents,
                    key=f"documents_{position}",
                    label_visibility="collapsed",
                )

                item["documents"] = selected_documents

                if selected_documents:

                    st.caption(
                        "Annexures will be assigned automatically "
                        "without gaps."
                    )

                    preview = []

                    for i, document in enumerate(
                        selected_documents
                    ):

                        preview.append(
                            f"**Annex-{chr(65 + i)}** — {document}"
                        )

                    st.markdown(
                        "\n\n".join(preview)
                    )

                else:

                    st.caption(
                        "No documents selected."
                    )


            # =================================================
            # INQUIRY COMMITTEE
            # =================================================

            elif base_name == "Inquiry Committee":

                st.markdown(
                    """
                    ### 👥 Inquiry Committee

                    Enter the details separately for each committee
                    role. Each role has its own clearly identified
                    group of fields.
                    """
                )

                committee_data = normalize_committee(
                    item.get("committee", {})
                )

                item["committee"] = committee_data

                for role_index, role in enumerate(
                    COMMITTEE_ROLES
                ):

                    st.markdown(
                        f"""
                        <div class="committee-role">
                            <div class="committee-role-title">
                                👤 {role}
                            </div>
                            <div class="committee-role-subtitle">
                                The following three fields belong
                                to this role.
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                    current = committee_data.get(
                        role,
                        {},
                    )

                    erp_key = (
                        f"committee_erp_"
                        f"{position}_{role_index}"
                    )

                    name_key = (
                        f"committee_name_"
                        f"{position}_{role_index}"
                    )

                    designation_key = (
                        f"committee_designation_"
                        f"{position}_{role_index}"
                    )

                    if erp_key not in st.session_state:
                        st.session_state[erp_key] = (
                            current.get("ERP#", "")
                        )

                    if name_key not in st.session_state:
                        st.session_state[name_key] = (
                            current.get("Name", "")
                        )

                    if designation_key not in st.session_state:
                        st.session_state[
                            designation_key
                        ] = current.get(
                            "Designation",
                            "",
                        )

                    c1, c2, c3 = st.columns(3)

                    with c1:

                        erp = st.text_input(
                            "ERP#",
                            key=erp_key,
                            placeholder="Enter ERP number",
                        )

                    with c2:

                        name = st.text_input(
                            "Name",
                            key=name_key,
                            placeholder="Enter name",
                        )

                    with c3:

                        designation = st.text_input(
                            "Designation",
                            key=designation_key,
                            placeholder="Enter designation",
                        )

                    committee_data[role] = {
                        "ERP#": erp,
                        "Name": name,
                        "Designation": designation,
                    }

                    st.markdown(
                        "<hr>",
                        unsafe_allow_html=True,
                    )

                item["committee"] = committee_data


            # =================================================
            # QUESTIONS / ANSWERS
            # =================================================

            elif base_name == (
                "Questions / Answers with the Accused"
            ):

                st.markdown(
                    """
                    ##### ❓ Questions & Answers

                    Add each question and its corresponding answer.
                    The generated document will present them as a
                    proper two-column table.
                    """
                )

                qa_rows = item.get(
                    "qa",
                    [],
                )

                qa_rows = list(qa_rows)

                if not qa_rows:
                    st.caption(
                        "No questions added yet."
                    )

                for q_index, row in enumerate(
                    qa_rows
                ):

                    q_col, a_col = st.columns(
                        [1, 1]
                    )

                    with q_col:

                        question = st.text_area(
                            f"Question {q_index + 1}",
                            value=row.get(
                                "question",
                                "",
                            ),
                            key=(
                                f"question_"
                                f"{position}_{q_index}"
                            ),
                            height=110,
                            placeholder=(
                                "Enter question..."
                            ),
                        )

                    with a_col:

                        answer = st.text_area(
                            f"Answer {q_index + 1}",
                            value=row.get(
                                "answer",
                                "",
                            ),
                            key=(
                                f"answer_"
                                f"{position}_{q_index}"
                            ),
                            height=110,
                            placeholder=(
                                "Enter answer..."
                            ),
                        )

                    qa_rows[q_index] = {
                        "question": question,
                        "answer": answer,
                    }

                item["qa"] = qa_rows

                if st.button(
                    "＋ Add Question / Answer",
                    key=f"add_qa_{position}",
                ):

                    qa_rows.append(
                        {
                            "question": "",
                            "answer": "",
                        }
                    )

                    item["qa"] = qa_rows

                    st.rerun()


            # =================================================
            # NORMAL TEXT / VOICE INDEX
            # =================================================

            else:

                text_key = (
                    f"inquiry_text_{position}"
                )

                voice_key = (
                    f"inquiry_voice_{position}"
                )

                voice_hash_key = (
                    f"inquiry_voice_hash_{position}"
                )

                if text_key not in st.session_state:

                    st.session_state[text_key] = str(
                        item.get(
                            "content",
                            "",
                        )
                        or ""
                    )

                st.markdown(
                    """
                    <div class="voice-card">
                        <div class="voice-title">
                            🎤 Voice + Text
                        </div>
                        <div class="voice-description">
                            Type your information or speak naturally.
                            Both methods use the same field.
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                audio = st.audio_input(
                    "🎤 Record information",
                    key=voice_key,
                )

                transcript = process_voice_input(
                    audio,
                    voice_hash_key,
                )

                if transcript:

                    existing = st.session_state.get(
                        text_key,
                        "",
                    ).strip()

                    if existing:

                        st.session_state[text_key] = (
                            existing
                            + "\n"
                            + transcript
                        )

                    else:

                        st.session_state[text_key] = (
                            transcript
                        )

                content = st.text_area(
                    "Information for this section",
                    key=text_key,
                    height=180,
                    placeholder=(
                        "Enter the relevant information "
                        "for this section..."
                    ),
                )

                item["content"] = content


            # =================================================
            # REMOVE SECTION
            # =================================================

            st.markdown("<br>", unsafe_allow_html=True)

            remove_section = st.button(
                "🗑️ Remove This Section",
                key=f"remove_section_{position}",
                use_container_width=True,
            )

            if remove_section:

                st.session_state.index_data.pop(
                    position
                )

                st.rerun()

            st.markdown(
                "</div>",
                unsafe_allow_html=True,
            )


        # ----------------------------------------------------
        # GENERATE INQUIRY
        # ----------------------------------------------------

        st.markdown("<br>", unsafe_allow_html=True)

        st.markdown(
            """
            <div class="section-card">
                <div class="section-title">
                    Ready to generate?
                </div>
                <div class="section-description">
                    DraftForge will include only the sections you
                    selected, in exactly the order you added them.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            '<div class="generate-button">',
            unsafe_allow_html=True,
        )

        generate_inquiry = st.button(
            "✦ Generate Inquiry Report",
            use_container_width=True,
            type="primary",
        )

        st.markdown(
            "</div>",
            unsafe_allow_html=True,
        )

        if generate_inquiry:

            if not st.session_state.index_data:

                st.warning(
                    "Please add at least one inquiry section."
                )

            else:

                with st.spinner(
                    "DraftForge is preparing your inquiry report..."
                ):

                    final_document = (
                        generate_inquiry_report(
                            st.session_state.index_data,
                            inquiry_type,
                        )
                    )

                st.session_state.generated_draft = (
                    final_document
                )

                st.session_state.editable_draft = (
                    final_document
                )

                st.session_state.editor_sync = (
                    final_document
                )

                save_history(
                    "E&D Inquiry",
                    final_document,
                )

                st.rerun()


    else:

        st.warning(
            "FFI Inquiry is currently under construction / under process."
        )


# ============================================================
# STEP 3 — GENERATE & EXPORT WORKSPACE
# ============================================================

if st.session_state.generated_draft:

    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown(
        """
        <div class="section-card">
            <div class="section-title">
                ③ Generate & Export
            </div>
            <div class="section-description">
                Review, edit and export your generated document.
                You remain in control of the final text.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


    # ========================================================
    # EDITOR STATE SYNCHRONIZATION
    # ========================================================

    if st.session_state.get(
        "editor_sync",
        "",
    ):

        st.session_state.document_editor = (
            st.session_state.editor_sync
        )

        st.session_state.editor_sync = ""

    elif not st.session_state.get(
        "document_editor",
        "",
    ):

        st.session_state.document_editor = (
            st.session_state.editable_draft
            or st.session_state.generated_draft
        )


    # ========================================================
    # DOCUMENT EDITOR
    # ========================================================

    st.markdown(
        """
        <div class="editor-card">
            <div class="editor-heading">
                📝 Document Editor
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    edited_document = st.text_area(
        "Document Editor",
        key="document_editor",
        height=550,
        label_visibility="collapsed",
    )


    # ========================================================
    # EDITOR ACTIONS
    # ========================================================

    edit_col1, edit_col2 = st.columns(
        [1, 1]
    )

    with edit_col1:

        if st.button(
            "💾 Save Changes",
            use_container_width=True,
        ):

            st.session_state.editable_draft = (
                edited_document
            )

            st.session_state.generated_draft = (
                edited_document
            )

            st.success(
                "Changes saved."
            )

    with edit_col2:

        if st.button(
            "↩️ Restore Original",
            use_container_width=True,
        ):

            original = (
                st.session_state.generated_draft
            )

            st.session_state.editable_draft = (
                original
            )

            st.session_state.editor_sync = (
                original
            )

            st.rerun()


    # ========================================================
    # AI EDITING ASSISTANT
    # ========================================================

    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown(
        """
        <div class="section-card">
            <div class="section-title">
                ✨ AI Editing Assistant
            </div>
            <div class="section-description">
                Ask DraftForge to improve, shorten, expand or correct
                the document while preserving its meaning.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


    if "edit_instruction_sync" in st.session_state:

        st.session_state.edit_instruction = (
            st.session_state.pop(
                "edit_instruction_sync"
            )
        )


    edit_instruction = st.text_area(
        "What would you like to change?",
        key="edit_instruction",
        height=120,
        placeholder=(
            "Example: Make this more concise and professional."
        ),
    )


    if st.button(
        "✨ Apply AI Editing",
        use_container_width=True,
    ):

        if not edit_instruction.strip():

            st.warning(
                "Please describe the change you want."
            )

        else:

            current_document = (
                st.session_state.document_editor
            )

            with st.spinner(
                "Applying your requested changes..."
            ):

                modified = ai_edit_document(
                    current_document,
                    edit_instruction,
                )

            if (
                modified.startswith(
                    "Groq API is not configured"
                )
                or modified.startswith(
                    "AI editing failed"
                )
            ):

                st.error(modified)

            else:

                st.session_state.editable_draft = (
                    modified
                )

                st.session_state.generated_draft = (
                    modified
                )

                st.session_state.editor_sync = (
                    modified
                )

                st.session_state.edit_instruction_sync = ""

                st.rerun()


    # ========================================================
    # EXPORTS
    # ========================================================

    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown(
        """
        <div class="export-card">
            <div class="section-title">
                📤 Export Document
            </div>
            <div class="section-description">
                Download your final document in the format you need.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    final_text = st.session_state.editable_draft

    export_cols = st.columns(4)


    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------

    with export_cols[0]:

        pdf_data = create_pdf(
            final_text
        )

        if pdf_data:

            st.download_button(
                "📕 PDF",
                data=pdf_data,
                file_name="draftforge_document.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

        else:

            st.warning(
                "PDF package unavailable."
            )


    # --------------------------------------------------------
    # DOCX
    # --------------------------------------------------------

    with export_cols[1]:

        docx_data = create_docx(
            final_text
        )

        if docx_data:

            st.download_button(
                "📘 DOCX",
                data=docx_data,
                file_name="draftforge_document.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                use_container_width=True,
            )

        else:

            st.warning(
                "DOCX package unavailable."
            )


    # --------------------------------------------------------
    # TXT
    # --------------------------------------------------------

    with export_cols[2]:

        st.download_button(
            "📄 TXT",
            data=final_text,
            file_name="draftforge_document.txt",
            mime="text/plain",
            use_container_width=True,
        )


    # --------------------------------------------------------
    # PNG
    # --------------------------------------------------------

    with export_cols[3]:

        png_data = create_png(
            final_text
        )

        if png_data:

            st.download_button(
                "🖼️ PNG",
                data=png_data,
                file_name="draftforge_document.png",
                mime="image/png",
                use_container_width=True,
            )

        else:

            st.warning(
                "PNG package unavailable."
            )


# ============================================================
# FOOTER
# ============================================================

st.markdown(
    """
    <div class="app-footer">
        ✦ DraftForge — AI Document Composer<br>
        Developed by: Raees Khan — Assistant Director, NADRA
    </div>
    """,
    unsafe_allow_html=True,
)
