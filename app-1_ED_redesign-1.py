import os
import io
import re
import json
import hashlib
import textwrap
from datetime import datetime

import streamlit as st

try:
    from groq import Groq
except Exception:
    Groq = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
except Exception:
    Document = None

try:
    from fpdf import FPDF
except Exception:
    FPDF = None

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None


st.set_page_config(
    page_title="DraftForge — AI Document Composer",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="expanded"
)


PROFILE_FILE = "user_profile.json"
HISTORY_FILE = "draftforge_history.json"

GROQ_MODEL = "openai/gpt-oss-120b"
WHISPER_MODEL = "whisper-large-v3"

DOCUMENT_TYPES = ["Email", "Letter", "Inquiry"]
INQUIRY_TYPES = ["E&D Inquiry", "FFI Inquiry"]

COMMITTEE_ROLES = [
    "Convener of Inquiry",
    "Member 1",
    "Member 2",
    "Departmental Representative"
]


DEFAULTS = {
    "document_type": "Email",
    "inquiry_type": "E&D Inquiry",
    "generated_draft": "",
    "editable_draft": "",
    "document_editor": "",
    "editor_sync": "",
    "edit_instruction": "",
    "edit_instruction_sync": "",
    "email_instruction": "",
    "letter_instruction": "",
    "email_voice_hash": "",
    "letter_voice_hash": "",
    "profile": {},
    "history": [],
    "show_history": False,
    "show_profile": False,
    "show_about": False,
    "ed_data": {},
    "ed_voice_hashes": {}
}


for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


st.markdown("""
<style>
.main .block-container {
    max-width: 1250px;
    padding-top: 1.5rem;
    padding-bottom: 4rem;
}

[data-testid="stSidebar"] {
    border-right: 1px solid rgba(128,128,128,.18);
}

.hero-box {
    padding: 1.5rem 1.7rem;
    border: 1px solid rgba(128,128,128,.20);
    border-radius: 18px;
    margin-bottom: 1.2rem;
    background: linear-gradient(
        135deg,
        rgba(120,120,120,.08),
        rgba(120,120,120,.025)
    );
}

.hero-title {
    font-size: 2rem;
    font-weight: 750;
    margin-bottom: .25rem;
}

.hero-subtitle {
    font-size: 1rem;
    opacity: .75;
}

.footer-note {
    text-align: center;
    opacity: .6;
    font-size: .82rem;
    margin-top: 2rem;
}

div[data-testid="stButton"] button,
div[data-testid="stDownloadButton"] button {
    border-radius: 10px;
    min-height: 42px;
}

textarea {
    font-size: 1rem !important;
}

.ed-card {
    padding: .8rem 1rem;
    border: 1px solid rgba(128,128,128,.18);
    border-radius: 12px;
    margin: .5rem 0;
    background: rgba(128,128,128,.025);
}
</style>
""", unsafe_allow_html=True)


# ---------------------------------------------------------
# JSON HELPERS
# ---------------------------------------------------------

def load_json(filename, default):
    try:
        if os.path.exists(filename):
            with open(filename, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass

    return default


def save_json(filename, data):
    try:
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(
                data,
                f,
                indent=2,
                ensure_ascii=False
            )
        return True
    except Exception:
        return False


# ---------------------------------------------------------
# PROFILE
# ---------------------------------------------------------

def profile_signature(profile):
    if not isinstance(profile, dict):
        return ""

    parts = []

    for label, key in [
        ("Name", "Name"),
        ("Designation", "Designation"),
        ("Contact No.", "Contact No."),
        ("Current Station", "Current Station")
    ]:
        val = str(profile.get(key, "") or "").strip()

        if val:
            if label in ("Name", "Designation"):
                parts.append(val)
            else:
                parts.append(f"{label}: {val}")

    return "\n".join(parts)


# ---------------------------------------------------------
# GROQ
# ---------------------------------------------------------

def get_groq_client():
    if Groq is None:
        return None

    key = None

    try:
        key = st.secrets.get("GROQ_API_KEY")
    except Exception:
        pass

    if not key:
        key = os.getenv("GROQ_API_KEY")

    try:
        return Groq(api_key=key) if key else None
    except Exception:
        return None


# ---------------------------------------------------------
# AI TEXT CLEANING
# ---------------------------------------------------------

def clean_ai_text(text):
    text = str(text or "").strip()

    text = re.sub(
        r"^```(?:text|markdown)?\s*",
        "",
        text,
        flags=re.I
    )

    text = re.sub(
        r"\s*```$",
        "",
        text
    )

    return text.strip()


# ---------------------------------------------------------
# VOICE TRANSCRIPTION
# ---------------------------------------------------------

def transcribe_audio(audio_bytes, filename="voice.wav"):
    client = get_groq_client()

    if client is None:
        st.error("Groq API key is not configured.")
        return ""

    try:
        f = io.BytesIO(audio_bytes)
        f.name = filename

        result = client.audio.transcriptions.create(
            file=(filename, f),
            model=WHISPER_MODEL,
            response_format="text"
        )

        if isinstance(result, str):
            return result.strip()

        return str(
            getattr(result, "text", result)
        ).strip()

    except Exception as e:
        st.error(f"Voice transcription failed: {e}")
        return ""


def process_voice_input(audio, state_key):
    if audio is None:
        return ""

    try:
        audio_bytes = audio.getvalue()
    except Exception:
        return ""

    if not audio_bytes:
        return ""

    audio_hash = hashlib.sha256(
        audio_bytes
    ).hexdigest()

    # Prevent the same recording being transcribed repeatedly
    if st.session_state.get(state_key) == audio_hash:
        return ""

    st.session_state[state_key] = audio_hash

    return transcribe_audio(audio_bytes)


# ---------------------------------------------------------
# EMAIL / LETTER GENERATION
# ---------------------------------------------------------

def generate_ai_document(
    document_type,
    instruction,
    profile
):
    client = get_groq_client()

    if client is None:
        raise RuntimeError(
            "Groq API key is missing. "
            "Add GROQ_API_KEY to Streamlit Secrets."
        )

    prompt = f"""
You are DraftForge, an AI assistant for
professional official correspondence.

Document type:
{document_type}

User's raw instructions:
{instruction}

Sender profile:
{profile_signature(profile)}

Prepare a professional official English document.

Correct grammar and speech-to-text errors.
Preserve the user's meaning.
Never invent facts.

Return only the document.
Do not add a sender signature.
"""

    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a professional official-document "
                    "drafting assistant. Never fabricate facts."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.2
    )

    return clean_ai_text(
        response.choices[0].message.content
    )


# ---------------------------------------------------------
# AI EDIT
# ---------------------------------------------------------

def ai_edit_document(document, instruction):
    client = get_groq_client()

    if client is None:
        raise RuntimeError(
            "Groq API key is missing."
        )

    prompt = f"""
Edit this official document according to the instruction.

Preserve meaning and facts.
Do not invent information.
Return only the revised document.

INSTRUCTION:
{instruction}

DOCUMENT:
{document}
"""

    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a careful professional "
                    "document editor. Never fabricate facts."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.1
    )

    return clean_ai_text(
        response.choices[0].message.content
    )


# ---------------------------------------------------------
# HISTORY
# ---------------------------------------------------------

def save_history(title, document):
    history = load_json(
        HISTORY_FILE,
        []
    )

    if not isinstance(history, list):
        history = []

    history.insert(
        0,
        {
            "title": title,
            "document": document,
            "date": datetime.now().strftime(
                "%Y-%m-%d %H:%M"
            )
        }
    )

    save_json(
        HISTORY_FILE,
        history[:30]
    )

    st.session_state.history = history[:30]


# ---------------------------------------------------------
# EXPORT
# ---------------------------------------------------------

def create_txt(text):
    return str(text).encode("utf-8")


def create_docx(text):
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed."
        )

    doc = Document()

    section = doc.sections[0]

    section.top_margin = Inches(.7)
    section.bottom_margin = Inches(.7)
    section.left_margin = Inches(.8)
    section.right_margin = Inches(.8)

    for line in str(text).splitlines():
        p = doc.add_paragraph()

        p.paragraph_format.space_after = Pt(6)

        if line.strip():
            r = p.add_run(line)
            r.font.name = "Arial"
            r.font.size = Pt(11)

    output = io.BytesIO()

    doc.save(output)

    output.seek(0)

    return output.getvalue()


def find_unicode_font():
    paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        os.path.join(
            os.getcwd(),
            "DejaVuSans.ttf"
        ),
        os.path.join(
            os.path.dirname(
                os.path.abspath(__file__)
            ),
            "DejaVuSans.ttf"
        ),
        r"C:\Windows\Fonts\arial.ttf",
        "/Library/Fonts/Arial.ttf"
    ]

    for path in paths:
        if os.path.isfile(path):
            return path

    for root in [
        "/usr/share/fonts",
        "/usr/local/share/fonts"
    ]:
        if os.path.isdir(root):
            for directory, _, files in os.walk(root):
                for filename in files:
                    if filename.lower() in (
                        "dejavusans.ttf",
                        "dejavusanscondensed.ttf"
                    ):
                        return os.path.join(
                            directory,
                            filename
                        )

    return None


def safe_pdf_text(text, width=88):
    output = []

    for raw in str(text).splitlines():

        if not raw.rstrip():
            output.append("")
            continue

        output.extend(
            textwrap.wrap(
                raw.rstrip(),
                width=width,
                break_long_words=True,
                break_on_hyphens=True,
                replace_whitespace=False,
                drop_whitespace=False
            ) or [""]
        )

    return output


def pdf_ascii_fallback(text):
    replacements = {
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2015": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2022": "-",
        "\u00a0": " ",
        "\u2026": "...",
        "\u00d7": "x",
        "\u2212": "-"
    }

    for old, new in replacements.items():
        text = str(text).replace(old, new)

    return text


def create_pdf(text):
    if FPDF is None:
        raise RuntimeError(
            "fpdf is not installed."
        )

    font = find_unicode_font()

    if font:
        try:
            pdf = FPDF()

            pdf.set_auto_page_break(
                auto=True,
                margin=15
            )

            pdf.add_page()

            try:
                pdf.add_font(
                    "DraftForgeUnicode",
                    "",
                    font
                )
            except TypeError:
                pdf.add_font(
                    "DraftForgeUnicode",
                    "",
                    font,
                    uni=True
                )

            pdf.set_font(
                "DraftForgeUnicode",
                size=11
            )

            width = (
                pdf.w -
                pdf.l_margin -
                pdf.r_margin
            )

            for line in safe_pdf_text(text):
                pdf.set_x(pdf.l_margin)

                if line == "":
                    pdf.ln(5)
                else:
                    pdf.multi_cell(
                        width,
                        6,
                        line,
                        border=0,
                        align="L"
                    )

            output = pdf.output()

            if isinstance(
                output,
                (bytearray, bytes)
            ):
                return bytes(output)

            return output.encode("latin-1")

        except Exception:
            pass

    pdf = FPDF()

    pdf.set_auto_page_break(
        auto=True,
        margin=15
    )

    pdf.add_page()

    pdf.set_font(
        "Helvetica",
        size=11
    )

    width = (
        pdf.w -
        pdf.l_margin -
        pdf.r_margin
    )

    fallback = pdf_ascii_fallback(text)

    for line in safe_pdf_text(fallback):
        pdf.set_x(pdf.l_margin)

        if line == "":
            pdf.ln(5)
        else:
            pdf.multi_cell(
                width,
                6,
                line.encode(
                    "latin-1",
                    "replace"
                ).decode("latin-1"),
                border=0,
                align="L"
            )

    output = pdf.output()

    if isinstance(
        output,
        (bytearray, bytes)
    ):
        return bytes(output)

    return output.encode("latin-1")


def create_png(text):
    if Image is None:
        raise RuntimeError(
            "Pillow is not installed."
        )

    try:
        font = ImageFont.truetype(
            "DejaVuSans.ttf",
            24
        )
    except Exception:
        font = ImageFont.load_default()

    lines = safe_pdf_text(
        text,
        65
    )

    image = Image.new(
        "RGB",
        (
            1400,
            max(
                400,
                120 + len(lines) * 38
            )
        ),
        "white"
    )

    draw = ImageDraw.Draw(image)

    y = 60

    for line in lines:
        draw.text(
            (60, y),
            line,
            fill="black",
            font=font
        )

        y += 38

    output = io.BytesIO()

    image.save(
        output,
        format="PNG"
    )

    output.seek(0)

    return output.getvalue()


# =========================================================
# E&D DATA
# =========================================================

ED_DEFAULTS = {
    "reference": "",
    "subject": "",
    "brief": "",
    "charges": [""],
    "accused_statement": "",

    "qa": [
        {
            "question": "",
            "answer": ""
        }
    ],

    "witnesses": [
        {
            "name": "",
            "designation": "",
            "statement": ""
        }
    ],

    "evidence": [
        {
            "category": "Departmental Evidence",
            "title": "",
            "description": ""
        }
    ],

    "cross_examination": "",
    "defence": "",
    "applicable_rules": "",

    # IMPORTANT:
    # Accused uses exactly the same three fields
    # as committee members.
    "accused": {
        "erp": "",
        "name": "",
        "designation": ""
    },

    "committee": {
        role: {
            "erp": "",
            "name": "",
            "designation": ""
        }
        for role in COMMITTEE_ROLES
    }
}


def init_ed():
    if not st.session_state.ed_data:
        st.session_state.ed_data = json.loads(
            json.dumps(ED_DEFAULTS)
        )


# =========================================================
# VOICE FIELD
# =========================================================

def field_voice(
    label,
    key,
    current,
    placeholder,
    height=150
):
    """
    Voice-enabled field.

    IMPORTANT FIX:
    The transcript is written into the widget's
    session-state key BEFORE the text_area is
    instantiated. This allows the transcript to
    survive Streamlit reruns and become part of
    the actual field value.
    """

    widget_key = f"text_{key}"

    audio = st.audio_input(
        f"🎤 Speak: {label}",
        key=f"audio_{key}"
    )

    transcript = process_voice_input(
        audio,
        f"voice_{key}"
    )

    if transcript:

        existing = st.session_state.get(
            widget_key,
            current or ""
        )

        existing = str(existing).strip()
        transcript = transcript.strip()

        if existing:
            combined = (
                existing +
                "\n" +
                transcript
            )
        else:
            combined = transcript

        combined = combined.strip()

        # Persist transcript BEFORE text_area
        st.session_state[widget_key] = combined

        current = combined

        st.success(
            f"Voice input added to {label}."
        )

    else:
        current = st.session_state.get(
            widget_key,
            current or ""
        )

    return st.text_area(
        label,
        key=widget_key,
        height=height,
        placeholder=placeholder
    )


# =========================================================
# COMMITTEE TEXT
# =========================================================

def committee_text(committee):
    lines = []

    for role in COMMITTEE_ROLES:

        member = committee.get(
            role,
            {}
        )

        values = []

        if member.get("erp"):
            values.append(
                f"ERP#: {member['erp']}"
            )

        if member.get("name"):
            values.append(
                f"Name: {member['name']}"
            )

        if member.get("designation"):
            values.append(
                f"Designation: {member['designation']}"
            )

        if values:
            lines.append(
                f"{role} — {', '.join(values)}"
            )

    return (
        "\n".join(lines)
        if lines
        else "No committee details provided."
    )


# =========================================================
# E&D REPORT GENERATION
# =========================================================

def generate_ed_report(data):

    client = get_groq_client()

    if client is None:
        raise RuntimeError(
            "Groq API key is missing. "
            "Add GROQ_API_KEY to Streamlit Secrets."
        )

    # -----------------------------------------------------
    # Evidence numbering
    # -----------------------------------------------------

    evidence = []

    annex_number = 0
    defence_number = 0

    for item in data["evidence"]:

        if (
            not item.get("title")
            and not item.get("description")
        ):
            continue

        if item.get("category") == "Defence Exhibit":

            defence_number += 1

            label = f"Ex-D{defence_number}"

        else:

            annex_number += 1

            label = f"Annex-{chr(64 + annex_number)}"

        evidence.append(
            {
                "label": label,
                "category": item.get(
                    "category",
                    ""
                ),
                "title": item.get(
                    "title",
                    ""
                ),
                "description": item.get(
                    "description",
                    ""
                )
            }
        )

    charges = [
        item.strip()
        for item in data["charges"]
        if item.strip()
    ]

    witnesses = [
        item
        for item in data["witnesses"]
        if any(
            str(value).strip()
            for value in item.values()
        )
    ]

    qa = [
        item
        for item in data["qa"]
        if (
            item.get("question", "").strip()
            or
            item.get("answer", "").strip()
        )
    ]

    # -----------------------------------------------------
    # COMPLETE PAYLOAD
    # -----------------------------------------------------

    payload = {
        "committee": data["committee"],

        # Explicit accused object
        "accused": data["accused"],

        "reference": data["reference"],
        "subject": data["subject"],
        "brief": data["brief"],
        "charges": charges,

        "accused_statement":
            data["accused_statement"],

        "qa": qa,
        "witnesses": witnesses,
        "evidence": evidence,

        "cross_examination":
            data["cross_examination"],

        "defence":
            data["defence"],

        "applicable_rules":
            data["applicable_rules"]
    }

    prompt = f"""
Prepare a concise, professionally structured
E&D inquiry report from the record below.

MANDATORY REPORT STRUCTURE AND NUMBERING:

1. Inquiry Committee / Composition
2. Inquiry Reference No.
3. Subject
4. Brief of the Inquiry
5. Articles of Charge / Allegations
6. Details and Statement of the Accused
7. Questions / Answers with the Accused
8. Statements of Witnesses / Officials
9. Documentary Evidence / Record Examined
10. Cross-Examination of Defence Witnesses
11. Defence / Written Explanation
12. Discussion, Assessment and Appreciation of Evidence
13. Findings on Each Charge
14. Conclusion
15. Specific Recommendations
16. Documents Recorded / Annexure Index
17. Inquiry Committee Signatures

RULES:

- Use ONLY facts in the supplied record.

- Never invent names, dates, events, evidence,
  admissions or rules.

- Keep the report concise but explain each issue
  specifically.

- Use bullets for witness particulars,
  evidence lists, charges and other subheadings
  where appropriate.

- Number charges independently.

- Assess every charge separately.

- In the Discussion section, weigh departmental
  and defence evidence.

- Assess credibility and reliability of witnesses.

- Consider corroboration, contradictions,
  admissions and documentary evidence.

- Findings must be based strictly on the evidence
  recorded and the balance of probabilities,
  not intuition.

- For each independent charge classify it exactly
  as:

  Proved

  Partially Proved

  Not Proved / Disproved

- Explain briefly why each classification follows
  from the evidence.

- Do not treat an allegation as proof merely
  because it was made.

- If evidence is insufficient or material
  contradictions remain unresolved, say so rather
  than filling gaps.

- Recommendations:

  If charges are not proved, recommend
  exoneration.

  If charges are proved or partially proved,
  recommend an appropriate statutory/proportionate
  penalty ONLY to the extent supported by the
  supplied applicable rules or penalty framework.

- Do not invent a rule or penalty provision.

- Defence documents must retain:
  Ex-D1, Ex-D2, Ex-D3 etc.

- Departmental documentary evidence must retain:
  Annex-A, Annex-B, Annex-C etc.

- The committee listed at the beginning must be
  repeated at the end for signatures.

- Under Section 6, explicitly report the accused's
  ERP#, Name and Designation from the supplied
  accused object.

- Do NOT say accused details are unavailable
  when any of these fields are supplied.

- The voice-transcribed text is part of the official
  record supplied by the user and must be reflected
  in the appropriate section.

- Do not add unsupported sections beyond the
  structure above.

RECORD:

{json.dumps(
    payload,
    ensure_ascii=False,
    indent=2
)}
"""

    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a senior official inquiry-report "
                    "drafting assistant. Evidence-based "
                    "reasoning only; never fabricate facts."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.1
    )

    return clean_ai_text(
        response.choices[0].message.content
    )


# =========================================================
# SIDEBAR
# =========================================================

with st.sidebar:

    st.markdown("## ✦ DraftForge")
    st.caption("AI Document Composer")

    st.divider()

    st.markdown("### Workspace")

    if st.button(
        "✦ New Document",
        use_container_width=True
    ):

        for key in [
            "generated_draft",
            "editable_draft",
            "document_editor",
            "editor_sync",
            "email_instruction",
            "letter_instruction"
        ]:
            st.session_state[key] = ""

        st.session_state.ed_data = {}

        st.session_state.document_type = "Email"
        st.session_state.inquiry_type = "E&D Inquiry"

        st.rerun()

    if st.button(
        "🗂 My Documents",
        use_container_width=True
    ):
        st.session_state.show_history = True
        st.session_state.show_profile = False
        st.session_state.show_about = False

    if st.button(
        "👤 My Profile",
        use_container_width=True
    ):
        st.session_state.show_profile = True
        st.session_state.show_history = False
        st.session_state.show_about = False

    if st.button(
        "ℹ About DraftForge",
        use_container_width=True
    ):
        st.session_state.show_about = True
        st.session_state.show_history = False
        st.session_state.show_profile = False

    st.divider()

    st.markdown("### 💡 Tips & Templates")

    st.info(
        "Write naturally; DraftForge converts "
        "your instructions into professional "
        "official English."
    )

    st.info(
        "🎤 Speak instead of typing."
    )

    st.info(
        "📑 E&D Inquiry uses a pre-designed "
        "evidence-based form; you no longer "
        "need to select indexes."
    )

    st.divider()

    st.caption(
        "Developed by: Raees Khan\n\n"
        "Assistant Director, NADRA"
    )


# =========================================================
# HERO
# =========================================================

st.markdown(
    """
    <div class='hero-box'>
        <div class='hero-title'>✦ DraftForge</div>
        <div class='hero-subtitle'>
        AI-powered workspace for professional official
        correspondence and inquiry documentation.
        </div>
    </div>
    """,
    unsafe_allow_html=True
)


# =========================================================
# PROFILE
# =========================================================

if st.session_state.show_profile:

    st.markdown("## 👤 My Profile")

    profile = load_json(
        PROFILE_FILE,
        {
            "Name": "",
            "Designation": "",
            "Contact No.": "",
            "Current Station": ""
        }
    )

    with st.form("profile_form"):

        name = st.text_input(
            "Name",
            value=profile.get(
                "Name",
                ""
            )
        )

        designation = st.text_input(
            "Designation",
            value=profile.get(
                "Designation",
                ""
            )
        )

        contact = st.text_input(
            "Contact No.",
            value=profile.get(
                "Contact No.",
                ""
            )
        )

        station = st.text_input(
            "Current Station",
            value=profile.get(
                "Current Station",
                ""
            )
        )

        if st.form_submit_button(
            "Save Profile",
            use_container_width=True
        ):

            new_profile = {
                "Name": name.strip(),
                "Designation": designation.strip(),
                "Contact No.": contact.strip(),
                "Current Station": station.strip()
            }

            if save_json(
                PROFILE_FILE,
                new_profile
            ):
                st.success(
                    "Profile saved successfully."
                )
            else:
                st.error(
                    "Unable to save profile."
                )


# =========================================================
# HISTORY
# =========================================================

if st.session_state.show_history:

    st.markdown("## 🗂 My Documents")

    history = load_json(
        HISTORY_FILE,
        []
    )

    if not history:
        st.info(
            "No saved documents yet."
        )

    for i, item in enumerate(history):

        with st.expander(
            f"{item.get('title', 'Untitled Document')} "
            f"— {item.get('date', '')}"
        ):

            saved = item.get(
                "document",
                ""
            )

            st.text_area(
                "Document",
                saved,
                height=250,
                key=f"history_{i}"
            )

            if st.button(
                "Restore",
                key=f"restore_{i}"
            ):

                st.session_state.generated_draft = saved
                st.session_state.editable_draft = saved
                st.session_state.editor_sync = saved
                st.session_state.show_history = False

                st.rerun()


# =========================================================
# ABOUT
# =========================================================

if st.session_state.show_about:

    st.markdown("## ℹ About DraftForge")

    st.write(
        "DraftForge is an AI-assisted drafting "
        "workspace for professional official "
        "correspondence and inquiry documentation."
    )

    st.markdown(
        """
        **Supported document types**

        - Email
        - Letter
        - E&D Inquiry
        - FFI Inquiry — Under Construction / Under Process
        """
    )


# =========================================================
# DOCUMENT TYPE
# =========================================================

st.markdown("## ① Choose Document")

st.caption(
    "Start by selecting what you want "
    "DraftForge to prepare."
)

columns = st.columns(3)

for i, document_type in enumerate(
    DOCUMENT_TYPES
):

    with columns[i]:

        selected = (
            "✓ "
            if st.session_state.document_type
            == document_type
            else ""
        )

        if st.button(
            selected + document_type,
            key=f"document_type_{document_type}",
            use_container_width=True
        ):

            st.session_state.document_type = (
                document_type
            )

            if document_type != "Inquiry":
                st.session_state.ed_data = {}

            st.session_state.generated_draft = ""
            st.session_state.editable_draft = ""
            st.session_state.editor_sync = ""

            st.rerun()


document_type = st.session_state.document_type


# =========================================================
# EMAIL / LETTER
# =========================================================

if document_type in (
    "Email",
    "Letter"
):

    st.markdown(
        "## ② Provide Information"
    )

    st.caption(
        "Describe what you want to communicate. "
        "You can type, speak, or use both."
    )

    profile = load_json(
        PROFILE_FILE,
        {}
    )

    prefix = document_type.lower()

    audio = st.audio_input(
        f"🎤 Speak your {prefix} instructions",
        key=f"{prefix}_audio"
    )

    transcript = process_voice_input(
        audio,
        f"{prefix}_voice_hash"
    )

    if transcript:

        current = st.session_state.get(
            f"{prefix}_instruction",
            ""
        ).strip()

        if current:
            st.session_state[
                f"{prefix}_instruction"
            ] = (
                current +
                "\n" +
                transcript
            ).strip()

        else:
            st.session_state[
                f"{prefix}_instruction"
            ] = transcript.strip()

        st.success(
            "Voice input added to the instruction box."
        )

    instruction = st.text_area(
        f"{document_type} Instructions",
        key=f"{prefix}_instruction",
        height=220,
        placeholder=(
            "Describe the purpose, recipient, "
            "facts and requested action in "
            "natural language."
        )
    )

    st.caption(
        "The microphone works even when "
        "the instruction box is empty."
    )

    if st.button(
        f"✦ Generate {document_type}",
        type="primary",
        use_container_width=True
    ):

        if not instruction.strip():

            st.warning(
                "Please provide information "
                "by typing or speaking."
            )

        else:

            with st.spinner(
                f"DraftForge is preparing "
                f"your {prefix}..."
            ):

                try:

                    document = generate_ai_document(
                        document_type,
                        instruction,
                        profile
                    )

                    signature = profile_signature(
                        profile
                    )

                    if signature:
                        document += (
                            "\n\n" +
                            signature
                        )

                    st.session_state.generated_draft = document
                    st.session_state.editable_draft = document
                    st.session_state.editor_sync = document

                    save_history(
                        document_type,
                        document
                    )

                    st.rerun()

                except Exception as e:
                    st.error(str(e))


# =========================================================
# E&D INQUIRY
# =========================================================

elif document_type == "Inquiry":

    st.markdown(
        "## ② E&D Inquiry — Structured Form"
    )

    inquiry_type = st.radio(
        "Inquiry Type",
        INQUIRY_TYPES,
        index=INQUIRY_TYPES.index(
            st.session_state.inquiry_type
        ),
        horizontal=True
    )

    st.session_state.inquiry_type = inquiry_type

    if inquiry_type == "FFI Inquiry":

        st.warning(
            "FFI Inquiry is currently "
            "Under Construction / Under Process."
        )

    else:

        init_ed()

        data = st.session_state.ed_data

        st.info(
            "Complete the pre-designed form below. "
            "Every section accepts natural-language "
            "text and voice input where applicable. "
            "DraftForge will intelligently number "
            "the final report and generate the "
            "evidence assessment, findings, "
            "conclusion and recommendations from "
            "the recorded material."
        )


        # -------------------------------------------------
        # 1 COMMITTEE
        # -------------------------------------------------

        st.markdown(
            "### 1. Inquiry Committee / Composition"
        )

        st.caption(
            "Enter the committee first. The same "
            "details will automatically appear again "
            "at the end for formal signatures."
        )

        for i, role in enumerate(
            COMMITTEE_ROLES
        ):

            member = data["committee"][role]

            st.markdown(
                f"**{role}**"
            )

            c1, c2, c3 = st.columns(3)

            with c1:
                member["erp"] = st.text_input(
                    "ERP#",
                    value=member.get(
                        "erp",
                        ""
                    ),
                    key=f"ed_erp_{i}"
                )

            with c2:
                member["name"] = st.text_input(
                    "Name",
                    value=member.get(
                        "name",
                        ""
                    ),
                    key=f"ed_name_{i}"
                )

            with c3:
                member["designation"] = st.text_input(
                    "Designation",
                    value=member.get(
                        "designation",
                        ""
                    ),
                    key=f"ed_des_{i}"
                )


        st.divider()


        # -------------------------------------------------
        # 2 REFERENCE
        # -------------------------------------------------

        st.markdown(
            "### 2. Inquiry Reference No."
        )

        data["reference"] = st.text_input(
            "Inquiry Reference No.",
            value=data["reference"],
            key="ed_reference",
            placeholder=(
                "Enter the official "
                "inquiry/reference number."
            )
        )


        # -------------------------------------------------
        # 3 SUBJECT
        # -------------------------------------------------

        st.markdown(
            "### 3. Subject"
        )

        data["subject"] = field_voice(
            "Subject",
            "subject",
            data["subject"],
            "State the subject of the inquiry in one clear line.",
            100
        )


        # -------------------------------------------------
        # 4 BRIEF
        # -------------------------------------------------

        st.markdown(
            "### 4. Brief of the Inquiry"
        )

        data["brief"] = field_voice(
            "Brief of the Inquiry",
            "brief",
            data["brief"],
            (
                "Briefly explain how the inquiry arose, "
                "relevant background, dates and circumstances. "
                "Include only known facts."
            )
        )


        # -------------------------------------------------
        # 5 CHARGES
        # -------------------------------------------------

        st.markdown(
            "### 5. Articles of Charge / Allegations"
        )

        st.caption(
            "Enter each independent charge separately. "
            "DraftForge will number and assess each "
            "charge independently."
        )

        for i in range(
            len(data["charges"])
        ):

            data["charges"][i] = field_voice(
                f"Charge {i + 1}",
                f"charge_{i}",
                data["charges"][i],
                (
                    "State one specific allegation/charge, "
                    "including relevant conduct, date, "
                    "place or rule only if known."
                ),
                120
            )


        c1, c2 = st.columns(2)

        with c1:

            if st.button(
                "＋ Add Another Charge",
                use_container_width=True
            ):

                data["charges"].append("")

                st.rerun()

        with c2:

            if (
                len(data["charges"]) > 1
                and
                st.button(
                    "− Remove Last Charge",
                    use_container_width=True
                )
            ):

                data["charges"].pop()

                st.rerun()


        # -------------------------------------------------
        # 6 ACCUSED
        # -------------------------------------------------

        st.markdown(
            "### 6. Details and Statement of the Accused"
        )

        st.caption(
            "Enter the accused's particulars using "
            "exactly the same fields used for each "
            "Inquiry Committee member: ERP#, Name "
            "and Designation."
        )

        accused = data["accused"]

        c1, c2, c3 = st.columns(3)

        with c1:

            accused["erp"] = st.text_input(
                "ERP#",
                value=accused.get(
                    "erp",
                    ""
                ),
                key="ed_accused_erp",
                placeholder="Accused ERP#"
            )

        with c2:

            accused["name"] = st.text_input(
                "Name",
                value=accused.get(
                    "name",
                    ""
                ),
                key="ed_accused_name",
                placeholder="Accused Name"
            )

        with c3:

            accused["designation"] = st.text_input(
                "Designation",
                value=accused.get(
                    "designation",
                    ""
                ),
                key="ed_accused_designation",
                placeholder="Accused Designation"
            )

        st.markdown(
            "**Statement of the Accused**"
        )

        data["accused_statement"] = field_voice(
            "Statement of the Accused",
            "accused_statement",
            data["accused_statement"],
            (
                "Record the accused person's "
                "explanation, admissions, denials "
                "and material assertions."
            )
        )


        # -------------------------------------------------
        # 7 QUESTIONS / ANSWERS
        # -------------------------------------------------

        st.markdown(
            "### 7. Questions / Answers with the Accused"
        )

        for i, row in enumerate(
            data["qa"]
        ):

            c1, c2 = st.columns(2)

            with c1:

                row["question"] = field_voice(
                    f"Question {i + 1}",
                    f"q_{i}",
                    row["question"],
                    "Enter the question put to the accused.",
                    100
                )

            with c2:

                row["answer"] = field_voice(
                    f"Answer {i + 1}",
                    f"a_{i}",
                    row["answer"],
                    "Enter the accused's answer.",
                    100
                )


        c1, c2 = st.columns(2)

        with c1:

            if st.button(
                "＋ Add Question / Answer",
                use_container_width=True
            ):

                data["qa"].append(
                    {
                        "question": "",
                        "answer": ""
                    }
                )

                st.rerun()

        with c2:

            if (
                len(data["qa"]) > 1
                and
                st.button(
                    "− Remove Last Q/A",
                    use_container_width=True
                )
            ):

                data["qa"].pop()

                st.rerun()


        # -------------------------------------------------
        # 8 WITNESSES
        # -------------------------------------------------

        st.markdown(
            "### 8. Statements of Witnesses / Officials"
        )

        for i, witness in enumerate(
            data["witnesses"]
        ):

            witness["name"] = st.text_input(
                "Witness / Official Name",
                value=witness.get(
                    "name",
                    ""
                ),
                key=f"wit_name_{i}",
                placeholder=(
                    "Name of witness/official"
                )
            )

            witness["designation"] = st.text_input(
                "Designation",
                value=witness.get(
                    "designation",
                    ""
                ),
                key=f"wit_des_{i}",
                placeholder=(
                    "Designation / office"
                )
            )

            witness["statement"] = field_voice(
                f"Statement of Witness {i + 1}",
                f"wit_stmt_{i}",
                witness.get(
                    "statement",
                    ""
                ),
                (
                    "Record the material statement. "
                    "Include facts personally known "
                    "to the witness."
                )
            )

            if i < len(
                data["witnesses"]
            ) - 1:
                st.divider()


        c1, c2 = st.columns(2)

        with c1:

            if st.button(
                "＋ Add Witness / Official",
                use_container_width=True
            ):

                data["witnesses"].append(
                    {
                        "name": "",
                        "designation": "",
                        "statement": ""
                    }
                )

                st.rerun()

        with c2:

            if (
                len(data["witnesses"]) > 1
                and
                st.button(
                    "− Remove Last Witness",
                    use_container_width=True
                )
            ):

                data["witnesses"].pop()

                st.rerun()


        # -------------------------------------------------
        # 9 EVIDENCE
        # -------------------------------------------------

        st.markdown(
            "### 9. Documentary Evidence / Record Examined"
        )

        st.caption(
            "Add departmental records and documents "
            "produced by the accused. Departmental "
            "evidence is automatically marked Annex-A, "
            "Annex-B... and defence documents Ex-D1, "
            "Ex-D2..."
        )

        for i, evidence in enumerate(
            data["evidence"]
        ):

            c1, c2 = st.columns(
                [1, 2]
            )

            with c1:

                evidence["category"] = st.selectbox(
                    "Evidence Category",
                    [
                        "Departmental Evidence",
                        "Defence Exhibit"
                    ],
                    index=(
                        0
                        if evidence.get(
                            "category"
                        ) != "Defence Exhibit"
                        else 1
                    ),
                    key=f"ev_cat_{i}"
                )

            with c2:

                evidence["title"] = st.text_input(
                    "Document / Record Title",
                    value=evidence.get(
                        "title",
                        ""
                    ),
                    key=f"ev_title_{i}",
                    placeholder=(
                        "e.g. Complaint/Application, "
                        "Official Record, Written Explanation"
                    )
                )

            evidence["description"] = field_voice(
                f"Evidence Details {i + 1}",
                f"ev_desc_{i}",
                evidence.get(
                    "description",
                    ""
                ),
                (
                    "Describe what the document contains "
                    "and why it is relevant. You may also "
                    "enter the file/document description."
                ),
                120
            )

            if i < len(
                data["evidence"]
            ) - 1:
                st.divider()


        c1, c2 = st.columns(2)

        with c1:

            if st.button(
                "＋ Add Evidence / Defence Exhibit",
                use_container_width=True
            ):

                data["evidence"].append(
                    {
                        "category":
                            "Departmental Evidence",
                        "title": "",
                        "description": ""
                    }
                )

                st.rerun()

        with c2:

            if (
                len(data["evidence"]) > 1
                and
                st.button(
                    "− Remove Last Evidence",
                    use_container_width=True
                )
            ):

                data["evidence"].pop()

                st.rerun()


        if st.button(
            "📎 Add Evidence Files (optional)",
            use_container_width=True
        ):
            st.session_state.show_file_uploader = True


        if st.session_state.get(
            "show_file_uploader",
            False
        ):

            uploaded_files = st.file_uploader(
                "Upload documentary evidence for reference/annexure naming",
                accept_multiple_files=True,
                key="ed_files"
            )

            if uploaded_files:

                st.session_state.ed_files = [
                    file.name
                    for file in uploaded_files
                ]

                st.caption(
                    "Files selected: "
                    +
                    ", ".join(
                        st.session_state.ed_files
                    )
                )


        # -------------------------------------------------
        # 10 CROSS EXAMINATION
        # -------------------------------------------------

        st.markdown(
            "### 10. Cross-Examination of Defence Witnesses"
        )

        data["cross_examination"] = field_voice(
            "Cross-Examination",
            "cross_examination",
            data["cross_examination"],
            (
                "Record the department's "
                "cross-examination of defence "
                "witnesses and material answers/"
                "admissions."
            )
        )


        # -------------------------------------------------
        # 11 DEFENCE
        # -------------------------------------------------

        st.markdown(
            "### 11. Defence / Written Explanation"
        )

        data["defence"] = field_voice(
            "Defence / Written Explanation",
            "defence",
            data["defence"],
            (
                "Record the written defence/"
                "explanation submitted by the "
                "accused and the material grounds "
                "relied upon."
            )
        )


        # -------------------------------------------------
        # RULES
        # -------------------------------------------------

        st.markdown(
            "### Applicable E&D Rules / Penalty Framework (Optional)"
        )

        data["applicable_rules"] = field_voice(
            "Applicable Rules / Penalty Framework",
            "rules",
            data["applicable_rules"],
            (
                "If known, provide the governing "
                "E&D rules or penalty framework "
                "so recommendations can be grounded "
                "correctly."
            ),
            120
        )


        # -------------------------------------------------
        # AI SECTIONS
        # -------------------------------------------------

        st.markdown(
            "### 12–15. AI Evidence Assessment, Findings, Conclusion & Recommendations"
        )

        st.info(
            "These sections are generated automatically. "
            "The AI will assess witness reliability and "
            "documentary evidence, map evidence to each "
            "charge, classify each charge as Proved / "
            "Partially Proved / Not Proved / Disproved, "
            "then formulate the conclusion and proportionate "
            "recommendations from the recorded evidence."
        )


        # -------------------------------------------------
        # 16 ANNEXURE INDEX
        # -------------------------------------------------

        st.markdown(
            "### 16. Documents Recorded / Annexure Index"
        )

        annex_number = 0
        defence_number = 0

        for evidence in data["evidence"]:

            if (
                not evidence.get("title")
                and
                not evidence.get("description")
            ):
                continue

            if (
                evidence.get("category")
                == "Defence Exhibit"
            ):

                defence_number += 1

                label = (
                    f"Ex-D{defence_number}"
                )

            else:

                annex_number += 1

                label = (
                    f"Annex-{chr(64 + annex_number)}"
                )

            st.write(
                f"**{label}** — "
                f"{evidence.get('title') or 'Untitled record'}"
            )

        st.caption(
            "The annexure/exhibit labels above are "
            "generated automatically and will be "
            "retained in the final report."
        )


        # -------------------------------------------------
        # 17 SIGNATURES
        # -------------------------------------------------

        st.markdown(
            "### 17. Inquiry Committee Signatures"
        )

        st.caption(
            "The committee details entered at the "
            "beginning will be repeated automatically "
            "here in the final report for signatures."
        )

        st.markdown(
            committee_text(
                data["committee"]
            )
        )

        st.divider()


        # -------------------------------------------------
        # GENERATE
        # -------------------------------------------------

        if st.button(
            "✦ Generate E&D Inquiry Report",
            type="primary",
            use_container_width=True
        ):

            if (
                not data["reference"].strip()
                and
                not data["subject"].strip()
                and
                not data["brief"].strip()
                and
                not any(
                    charge.strip()
                    for charge in data["charges"]
                )
            ):

                st.warning(
                    "Please provide the inquiry "
                    "reference/subject/brief and "
                    "at least one charge before "
                    "generating."
                )

            else:

                with st.spinner(
                    "DraftForge is assessing the "
                    "recorded evidence and preparing "
                    "the E&D inquiry report..."
                ):

                    try:

                        report = generate_ed_report(
                            data
                        )

                        st.session_state.generated_draft = report
                        st.session_state.editable_draft = report
                        st.session_state.editor_sync = report

                        save_history(
                            "E&D Inquiry",
                            report
                        )

                        st.rerun()

                    except Exception as e:
                        st.error(str(e))


# =========================================================
# DOCUMENT WORKSPACE
# =========================================================

if st.session_state.get(
    "generated_draft",
    ""
).strip():

    st.divider()

    st.markdown(
        "## ✎ Document Workspace"
    )

    st.caption(
        "Review the generated document, make "
        "manual changes, or ask DraftForge "
        "to edit it."
    )

    if st.session_state.get(
        "editor_sync",
        ""
    ):

        st.session_state.document_editor = (
            st.session_state.editor_sync
        )

        st.session_state.editor_sync = ""

    elif not st.session_state.get(
        "document_editor",
        ""
    ):

        st.session_state.document_editor = (
            st.session_state.editable_draft
            or
            st.session_state.generated_draft
        )


    edited = st.text_area(
        "Document",
        key="document_editor",
        height=520
    )


    c1, c2, c3 = st.columns(3)

    with c1:

        if st.button(
            "💾 Save Changes",
            use_container_width=True
        ):

            st.session_state.editable_draft = edited
            st.session_state.generated_draft = edited

            save_history(
                "Edited Document",
                edited
            )

            st.success(
                "Changes saved."
            )


    with c2:

        if st.button(
            "↩ Restore Original",
            use_container_width=True
        ):

            st.session_state.editor_sync = (
                st.session_state.generated_draft
            )

            st.rerun()


    with c3:

        st.download_button(
            "⬇ Download TXT",
            data=create_txt(edited),
            file_name="DraftForge_Document.txt",
            mime="text/plain",
            use_container_width=True
        )


    # -----------------------------------------------------
    # AI EDIT
    # -----------------------------------------------------

    st.markdown(
        "### ✨ AI Edit"
    )

    if st.session_state.get(
        "edit_instruction_sync",
        ""
    ):

        st.session_state.edit_instruction = (
            st.session_state.edit_instruction_sync
        )

        st.session_state.edit_instruction_sync = ""


    instruction = st.text_area(
        "Tell DraftForge what you want changed",
        key="edit_instruction",
        height=120,
        placeholder=(
            "Example: Make this more concise and formal."
        )
    )


    if st.button(
        "✨ Apply AI Edit",
        use_container_width=True
    ):

        if not instruction.strip():

            st.warning(
                "Please enter an editing instruction."
            )

        else:

            with st.spinner(
                "DraftForge is editing your document..."
            ):

                try:

                    edited_document = ai_edit_document(
                        edited,
                        instruction
                    )

                    st.session_state.editable_draft = edited_document
                    st.session_state.generated_draft = edited_document
                    st.session_state.editor_sync = edited_document

                    st.rerun()

                except Exception as e:
                    st.error(str(e))


    # -----------------------------------------------------
    # EXPORT
    # -----------------------------------------------------

    st.divider()

    st.markdown(
        "### 📤 Export"
    )

    c1, c2, c3 = st.columns(3)

    with c1:

        try:

            st.download_button(
                "📄 Download DOCX",
                data=create_docx(edited),
                file_name="DraftForge_Document.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                use_container_width=True
            )

        except Exception as e:

            st.error(
                f"DOCX export unavailable: {e}"
            )


    with c2:

        try:

            st.download_button(
                "📕 Download PDF",
                data=create_pdf(edited),
                file_name="DraftForge_Document.pdf",
                mime="application/pdf",
                use_container_width=True
            )

        except Exception as e:

            st.error(
                f"PDF export unavailable: {e}"
            )


    with c3:

        try:

            st.download_button(
                "🖼 Download PNG",
                data=create_png(edited),
                file_name="DraftForge_Document.png",
                mime="image/png",
                use_container_width=True
            )

        except Exception as e:

            st.error(
                f"PNG export unavailable: {e}"
            )


# =========================================================
# FOOTER
# =========================================================

st.markdown(
    """
    <div class='footer-note'>
    DraftForge — AI Document Composer<br>
    Developed by: Raees Khan — Assistant Director, NADRA
    </div>
    """,
    unsafe_allow_html=True
    )
